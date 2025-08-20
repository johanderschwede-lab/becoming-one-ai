# build this

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_manual():
    """Create nicely formatted Word manual"""
    doc = Document()
    
    # Document styling
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # TODO: Add logo once available
    # doc.add_picture('compass_logo.png', width=Inches(2.5))
    
    # Title
    title = doc.add_heading('Compass Builder User Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('A friendly guide to managing your wisdom and method content')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = doc.styles['Subtitle']
    
    # Add a page break after the cover
    doc.add_page_break()
    
    # Read content
    with open('Marianne_User_Manual.md', 'r') as f:
        content = f.read()
    
    # Remove the title and subtitle as we've already added them
    content = content.split('---')[0]  # Remove everything after the last horizontal rule
    content = '\n'.join(content.split('\n')[3:])  # Skip the original title and subtitle
    
    # Process each section
    sections = content.split('\n## ')
    
    for section in sections:
        if not section.strip():
            continue
            
        # Section title
        title_end = section.find('\n')
        title = section[:title_end].strip('# ')  # Remove any remaining #
        content = section[title_end:].strip()
        
        # Add section heading
        heading = doc.add_heading(title, 1)
        heading.style.font.color.rgb = RGBColor(30, 136, 229)  # Blue
        
        # Process subsections
        subsections = content.split('\n### ')
        
        for i, subsection in enumerate(subsections):
            if not subsection.strip():
                continue
                
            if i == 0 and not subsection.startswith('#'):
                # This is the main section content
                for para in subsection.split('\n\n'):
                    if para.strip():
                        if para.startswith('- '):
                            # Bullet points
                            for line in para.split('\n'):
                                if line.strip():
                                    p = doc.add_paragraph(line[2:], style='List Bullet')
                        elif para.startswith('1. '):
                            # Numbered list
                            for line in para.split('\n'):
                                if line.strip():
                                    num_match = re.match(r'(\d+)\. (.*)', line)
                                    if num_match:
                                        p = doc.add_paragraph(num_match.group(2), style='List Number')
                        else:
                            p = doc.add_paragraph(para)
                continue
            
            # Subsection title
            sub_title_end = subsection.find('\n')
            sub_title = subsection[:sub_title_end].strip('# ')
            sub_content = subsection[sub_title_end:].strip()
            
            # Add subsection heading
            heading = doc.add_heading(sub_title, 2)
            heading.style.font.color.rgb = RGBColor(66, 66, 66)  # Dark gray
            
            # Process paragraphs
            for para in sub_content.split('\n\n'):
                if not para.strip():
                    continue
                    
                if para.startswith('- '):
                    # Bullet points
                    for line in para.split('\n'):
                        if line.strip():
                            p = doc.add_paragraph(line[2:], style='List Bullet')
                elif para.startswith('1. '):
                    # Numbered list
                    for line in para.split('\n'):
                        if line.strip():
                            num_match = re.match(r'(\d+)\. (.*)', line)
                            if num_match:
                                p = doc.add_paragraph(num_match.group(2), style='List Number')
                else:
                    p = doc.add_paragraph(para)
                
                # Add some spacing
                p.paragraph_format.space_after = Pt(12)
    
    # Save the document
    doc.save('Compass Builder User Guide.docx')
    print("Created 'Compass Builder User Guide.docx'")

if __name__ == '__main__':
    create_manual()